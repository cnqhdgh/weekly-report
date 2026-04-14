"""
AI 주간 보고서 생성기 — Streamlit + Gemini + python-docx
"""

from __future__ import annotations

import io
import json
import re
from copy import deepcopy
from datetime import date, datetime
from typing import Any

from google import genai
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
try:
    from hwpx import HwpxDocument
except Exception:
    HwpxDocument = None


def get_api_key() -> str | None:
    try:
        k = st.secrets["GEMINI_API_KEY"]
        if str(k).strip():
            return str(k).strip()
    except Exception:
        return None
    return None


def format_item_block(
    title: str,
    item_date: str | None,
    location: str,
    people: str,
    content: str,
    *,
    done_suffix: bool = False,
    progress_suffix: bool = False,
) -> str:
    sublines: list[str] = []
    if item_date:
        sublines.append(f" - 일자: {item_date}")
    if location:
        sublines.append(f"  · 장소: {location}")
    if people:
        sublines.append(f"  · 인원: {people}")
    if content:
        sublines.append(f"  · 내용: {content}")

    title = (title or "").strip()
    if title:
        head = f"○ {title}"
        body = f"{head}\n" + "\n".join(sublines) if sublines else head
    elif sublines:
        body = "\n".join(sublines)
    else:
        body = "(내용 없음)"

    if done_suffix:
        return f"{body}\n[완료]"
    if progress_suffix:
        return f"{body}\n[진행중]"
    return body


def _parse_iso_date(s: str) -> date | None:
    s = (s or "").strip()[:10]
    if len(s) < 10:
        return None
    try:
        return date.fromisoformat(s)
    except ValueError:
        return None


def _memo_has_paren_complete(*parts: str) -> bool:
    """메모에 (완)/[완] 표시가 있는지(_strip_done_markers와 동일 규칙). '완료' 단어만으로는 False."""
    blob = " ".join(p for p in parts if p).replace("（", "(").replace("）", ")")
    return bool(
        re.search(r"\(\s*완\s*\)", blob) or re.search(r"\[\s*완\s*\]", blob)
    )


def _strip_done_markers(text: str) -> str:
    """표시용 텍스트에서 완료 마커를 제거한다."""
    if not text:
        return ""
    cleaned = re.sub(r"\(\s*완\s*\)|\[\s*완\s*\]", "", text)
    return re.sub(r"\s{2,}", " ", cleaned).strip()


def reconcile_category(it: dict[str, Any], ref: date) -> str:
    """date가 없거나 기준일과 같을 때만 사용. (이후/이전은 build_table_rows에서 먼저 처리)"""
    d = _parse_iso_date(str(it.get("date") or ""))
    if d is not None and d > ref:
        return "current_plan"

    blob = " ".join(
        [str(it.get(k) or "") for k in ("title", "content", "location", "people")]
        + [str(it.get("date") or "").strip()]
    ).lower()
    if (
        _memo_has_paren_complete(
            str(it.get("title") or ""),
            str(it.get("content") or ""),
            str(it.get("location") or ""),
            str(it.get("people") or ""),
            str(it.get("date") or "").strip(),
        )
        or "완료" in blob
    ):
        return "prior_done"
    if d is not None:
        return "prior_plan" if d < ref else "current_plan"
    cat = (it.get("category") or "").strip().lower()
    if cat == "prior_done":
        return "prior_done"
    if cat == "prior_plan":
        return "prior_plan"
    return "current_plan"


def build_table_rows(
    items: list[dict[str, Any]],
    ref: date,
) -> list[dict[str, str]]:
    """Gemini 결과를 화면 표시용 3열 행으로 변환."""
    rows: list[dict[str, str]] = []

    for it in items:
        title_raw = str(it.get("title") or "").strip()
        item_date = it.get("date")
        if item_date is not None and str(item_date).strip():
            d_raw = str(item_date).strip()[:10]
        else:
            d_raw = ""
        location_raw = str(it.get("location") or "").strip()
        people_raw = str(it.get("people") or "").strip()
        content_raw = str(it.get("content") or "").strip()

        date_raw = str(item_date).strip() if item_date is not None else ""
        has_paren_done = _memo_has_paren_complete(
            title_raw, content_raw, location_raw, people_raw, date_raw
        )

        # "(완)"은 분류에만 사용하고 표 본문에는 노출하지 않음
        title = _strip_done_markers(title_raw)
        location = _strip_done_markers(location_raw)
        people = _strip_done_markers(people_raw)
        content = _strip_done_markers(content_raw)

        d_parsed = _parse_iso_date(d_raw) if d_raw else None
        base = format_item_block(title, d_raw or None, location, people, content)
        done = format_item_block(title, d_raw or None, location, people, content, done_suffix=True)
        in_progress = format_item_block(
            title, d_raw or None, location, people, content, progress_suffix=True
        )

        if d_parsed is not None and d_parsed > ref:
            rows.append({"전주계획": "", "전주실적": "", "금주계획": base})
        elif d_parsed is not None and d_parsed < ref:
            if has_paren_done:
                rows.append({"전주계획": base, "전주실적": done, "금주계획": ""})
            else:
                rows.append({"전주계획": base, "전주실적": in_progress, "금주계획": ""})
        else:
            cat = reconcile_category(it, ref)
            if cat == "prior_done":
                rows.append({"전주계획": base, "전주실적": done, "금주계획": ""})
            elif cat == "prior_plan":
                rows.append(
                    {
                        "전주계획": base,
                        "전주실적": done if has_paren_done else in_progress,
                        "금주계획": "",
                    }
                )
            else:
                rows.append({"전주계획": "", "전주실적": "", "금주계획": base})

    if not rows:
        rows.append({"전주계획": "", "전주실적": "", "금주계획": ""})
    return rows


def align_three_columns(
    a: list[str], b: list[str], c: list[str]
) -> list[dict[str, str]]:
    n = max(len(a), len(b), len(c), 1)
    a = a + [""] * (n - len(a))
    b = b + [""] * (n - len(b))
    c = c + [""] * (n - len(c))
    rows = []
    for i in range(n):
        rows.append({"전주계획": a[i], "전주실적": b[i], "금주계획": c[i]})
    return rows


def table_to_plain_text(rows: list[dict[str, str]]) -> str:
    lines = ["전주계획\t전주실적\t금주계획"]
    for r in rows:
        lines.append(
            "\t".join(
                [
                    (r.get("전주계획") or "").replace("\n", " "),
                    (r.get("전주실적") or "").replace("\n", " "),
                    (r.get("금주계획") or "").replace("\n", " "),
                ]
            )
        )
    return "\n".join(lines)


def render_report_table(rows: list[dict[str, str]]) -> None:
    def esc_html(s: str) -> str:
        return (
            s.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
        )

    table_rows = []
    for r in rows:
        col1 = esc_html(r.get("전주계획") or "").replace("\n", "<br>")
        col2 = esc_html(r.get("전주실적") or "").replace("\n", "<br>")
        col3 = esc_html(r.get("금주계획") or "").replace("\n", "<br>")
        table_rows.append(
            "<tr>"
            f"<td>{col1}</td>"
            f"<td>{col2}</td>"
            f"<td>{col3}</td>"
            "</tr>"
        )

    html = f"""
<style>
.weekly-report-wrap {{
  border: 1px solid #e5e7eb;
  border-radius: 12px;
  overflow: hidden;
  background: #ffffff;
}}
.weekly-report-table {{
  width: 100%;
  border-collapse: collapse;
  table-layout: fixed;
}}
.weekly-report-table th {{
  background: linear-gradient(180deg, #f8fafc 0%, #eef2ff 100%);
  color: #111827;
  font-weight: 700;
  font-size: 0.95rem;
  text-align: center;
  padding: 12px;
  border-bottom: 1px solid #e5e7eb;
  border-right: 1px solid #e5e7eb;
}}
.weekly-report-table th:last-child {{
  border-right: none;
}}
.weekly-report-table td {{
  vertical-align: top;
  padding: 12px 14px;
  font-size: 0.9rem;
  line-height: 1.55;
  color: #1f2937;
  border-bottom: 1px solid #f1f5f9;
  border-right: 1px solid #f1f5f9;
  white-space: normal;
  word-break: break-word;
}}
.weekly-report-table td:last-child {{
  border-right: none;
}}
.weekly-report-table tr:last-child td {{
  border-bottom: none;
}}
</style>
<div class="weekly-report-wrap">
  <table class="weekly-report-table">
    <thead>
      <tr>
        <th>전주계획</th>
        <th>전주실적</th>
        <th>금주계획</th>
      </tr>
    </thead>
    <tbody>
      {''.join(table_rows)}
    </tbody>
  </table>
</div>
"""
    components.html(html, height=430, scrolling=True)


def _set_cell_multiline(cell: Any, text: str, font_pt: int = 10) -> None:
    lines = text.split("\n") if text else [""]
    cell.text = lines[0]
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(font_pt)
    for line in lines[1:]:
        p = cell.add_paragraph(line)
        for r in p.runs:
            r.font.size = Pt(font_pt)


def build_docx(rows: list[dict[str, str]]) -> bytes:
    doc = Document()
    title = doc.add_paragraph("AI 주간 보고서")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    doc.add_paragraph(f"작성 시각: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["전주계획", "전주실적", "금주계획"]):
        p = hdr[i].paragraphs[0]
        p.text = h
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r in rows:
        row_cells = table.add_row().cells
        vals = [r.get("전주계획") or "", r.get("전주실적") or "", r.get("금주계획") or ""]
        for idx, text in enumerate(vals):
            _set_cell_multiline(row_cells[idx], text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_hwpx(rows: list[dict[str, str]]) -> bytes:
    """python-hwpx를 사용해 한컴 호환 HWPX 생성."""
    if HwpxDocument is None:
        raise RuntimeError("python-hwpx 라이브러리를 찾을 수 없습니다.")
    doc = HwpxDocument.new()
    # 한글 문서를 가로 방향으로 설정하고 여백을 줄여 표를 넓게 사용
    sec = doc.sections[0]
    sec.properties.set_page_size(width=84186, height=59528, orientation="WIDELY")
    margin_20mm = 5669
    sec.properties.set_page_margins(
        left=margin_20mm,
        right=margin_20mm,
        top=margin_20mm,
        bottom=margin_20mm,
        header=0,
        footer=0,
        gutter=0,
    )

    doc.add_paragraph("AI 주간 보고서")
    doc.add_paragraph(f"작성 시각: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    content_width = 84186 - (margin_20mm * 2)
    table = doc.add_table(rows=max(len(rows), 1) + 1, cols=3, width=content_width)
    table.set_cell_text(0, 0, "전주계획")
    table.set_cell_text(0, 1, "전주실적")
    table.set_cell_text(0, 2, "금주계획")
    # 헤더 행 가운데 정렬 전용 문단 속성 추가/적용
    center_prop_id = "999"
    if center_prop_id not in doc.paragraph_properties:
        base = deepcopy(doc.paragraph_properties["0"])
        base.id = int(center_prop_id)
        base.raw_id = center_prop_id
        base.align.horizontal = "CENTER"
        base.align.attributes["horizontal"] = "CENTER"
        doc.paragraph_properties[center_prop_id] = base
    for c in range(3):
        header_para = table.cell(0, c).paragraphs[0]
        header_para.para_pr_id_ref = center_prop_id

    for idx, row in enumerate(rows, start=1):
        table.set_cell_text(idx, 0, row.get("전주계획") or "")
        table.set_cell_text(idx, 1, row.get("전주실적") or "")
        table.set_cell_text(idx, 2, row.get("금주계획") or "")

    buf = io.BytesIO()
    doc.save_to_stream(buf)
    return buf.getvalue()


def call_gemini(memo: str, ref: date) -> list[dict[str, Any]]:
    client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
    primary_model = "gemini-2.0-flash"
    fallback_model = "gemini-3-flash-preview"

    ref_s = ref.isoformat()
    system = f"""당신은 한국어 업무 메모를 구조화하는 비서입니다.
기준 날짜(reference_date): {ref_s}

각 메모에서 활동/일정 단위로 항목을 나누고, 다음 필드를 추출하세요:
- title: 제목(한 줄 요약)
- date: YYYY-MM-DD 형식. 메모에 명시된 날짜만 사용. 없으면 null
- location: 장소(없으면 빈 문자열)
- people: 인원/참석자(없으면 빈 문자열)
- content: 상세 내용
- category: 아래 중 하나만 사용
  * prior_done: 해당 항목에 '(완)' 표시가 **있을 때만**(전주실적). '(완)'이 없으면 prior_done을 쓰지 마세요.
  * prior_plan: '(완)'이 없고, date가 기준 날짜 **미만**인 계획/예정
  * current_plan: '(완)'이 없고, date가 기준 날짜 **이상**이거나 date가 null인 예정 항목

분류 규칙(중요, 앱에서 date로 최종 배치함):
1) date가 기준일보다 **이후**이면 (완) 여부와 관계없이 금주계획 열만 사용.
2) date가 기준일보다 **이전**이면: 전주계획·전주실적에 동일 본문을 넣고, '(완)' 있으면 전주실적에 [완료], 없으면 [진행중](앱 처리).
3) date가 없거나 기준일과 같을 때: prior_done은 위와 동일하게 전주계획·전주실적([완료] 등). prior_plan도 전주계획·전주실적에 넣고 실적 열은 (완) 유무로 [완료]/[진행중].

반드시 JSON만 출력:
{{"items":[{{"title":"...","date":"YYYY-MM-DD 또는 null","location":"...","people":"...","content":"...","category":"prior_done|prior_plan|current_plan"}}]}}
"""
    prompt = f"{system}\n\n메모:\n{memo}"
    resp = None
    first_error: Exception | None = None
    for model_name in (primary_model, fallback_model):
        try:
            resp = client.models.generate_content(
                model=model_name,
                contents=prompt,
                config={
                    "temperature": 0.2,
                    "response_mime_type": "application/json",
                },
            )
            break
        except Exception as e:
            if first_error is None:
                first_error = e
            # 요청대로 2.0이 실패하면 3-flash-preview 자동 재시도
            continue

    if resp is None:
        message = str(first_error) if first_error else "알 수 없는 오류"
        if "404" in message or "NOT_FOUND" in message:
            try:
                available = []
                for m in client.models.list():
                    name = getattr(m, "name", "")
                    if "gemini" in name:
                        available.append(name)
                sample = ", ".join(available[:20]) if available else "(조회 결과 없음)"
                raise RuntimeError(
                    f"모델 호출 실패. 1차: {primary_model}, 2차: {fallback_model}. "
                    f"사용 가능한 모델 예시: {sample}"
                ) from first_error
            except Exception as list_err:
                raise RuntimeError(
                    f"모델 호출 실패. 1차: {primary_model}, 2차: {fallback_model}. "
                    f"모델 목록 조회도 실패했습니다: {list_err}"
                ) from first_error
        raise RuntimeError(
            f"모델 호출 실패. 1차: {primary_model}, 2차: {fallback_model}. 원인: {message}"
        ) from first_error
    text = (resp.text or "").strip()
    data = json.loads(text)
    items = data.get("items") or []
    if not isinstance(items, list):
        return []
    return [x for x in items if isinstance(x, dict)]


def main() -> None:
    st.set_page_config(page_title="AI 주간 보고서 생성기", layout="wide")
    st.title("AI 주간 보고서 생성기")

    with st.expander("사용 방법 가이드", expanded=False):
        st.info(
            """
**AI 자동 분류:** 메모 형식으로 내용을 자유롭게 입력하면 AI가 제목, 날짜, 장소, 인원, 내용을 자동으로 분류하여 보고서를 작성합니다.

**주간 자동 배정:** 보고서 생성 날짜를 기준으로 '전주계획', '전주실적', '금주계획'을 스마트하게 분류해 표로 만듭니다.

**전주 표시:** 보고서 기준일보다 **이전 날짜**인 항목은 전주계획·전주실적에 모두 표시됩니다. **(완)**이 있으면 전주실적 끝에 **[완료]**, 없으면 **[진행중]**이 붙습니다.
            """.strip(),
        )

    api_key = get_api_key()
    if not api_key:
        st.warning(
            "Gemini API 키가 필요합니다. `.streamlit/secrets.toml`에 "
            "`GEMINI_API_KEY = \"...\"` 를 추가하세요."
        )

    memo = st.text_area(
        "비정형 메모 입력",
        height=280,
        placeholder="예: 4/7 팀 회의 R&D센터 (완)\n4/10 고객사 방문 예정 ...",
    )

    default_today = date.today()
    ref_date = st.date_input("보고서 기준 날짜", value=default_today)

    if st.button("보고서 생성", type="primary", disabled=not api_key or not memo.strip()):
        with st.spinner("Gemini가 메모를 분석하는 중입니다…"):
            try:
                raw_items = call_gemini(memo.strip(), ref_date)
            except Exception as e:
                st.error(f"Gemini 호출 또는 JSON 파싱 오류: {e}")
                return

        rows = build_table_rows(raw_items, ref_date)
        st.session_state["report_rows"] = rows
        st.session_state["report_plain"] = table_to_plain_text(rows)
        st.session_state["report_docx"] = build_docx(rows)
        try:
            st.session_state["report_hwpx"] = build_hwpx(rows)
            st.session_state["report_hwpx_error"] = ""
        except Exception as e:
            st.session_state["report_hwpx"] = b""
            st.session_state["report_hwpx_error"] = str(e)

    rows = st.session_state.get("report_rows")
    if rows:
        st.subheader("생성 결과")
        render_report_table(rows)

        plain = st.session_state.get("report_plain", "")
        safe = json.dumps(plain)
        components.html(
            f"""
<div style="font-family:sans-serif;margin:0 0 8px 0;">
  <button id="cpbtn" style="padding:0.4rem 0.9rem;cursor:pointer;border-radius:6px;border:1px solid #ccc;background:#f7f7f7;">
    텍스트 복사
  </button>
  <span id="cpmsg" style="margin-left:10px;color:#2e7d32;font-size:0.9rem;"></span>
</div>
<script>
  const t = {safe};
  const btn = document.getElementById('cpbtn');
  const msg = document.getElementById('cpmsg');
  btn.addEventListener('click', async () => {{
    try {{
      await navigator.clipboard.writeText(t);
      msg.textContent = '클립보드에 복사했습니다.';
    }} catch (e) {{
      msg.textContent = '복사에 실패했습니다. 브라우저 권한을 확인하세요.';
      msg.style.color = '#c62828';
    }}
  }});
</script>
""",
            height=70,
        )

        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button(
                label="텍스트 파일 다운로드 (.txt)",
                data=plain.encode("utf-8-sig"),
                file_name="weekly_report.txt",
                mime="text/plain; charset=utf-8",
            )
        docx_bytes = st.session_state.get("report_docx")
        with c2:
            st.download_button(
                label="워드(.docx) 파일 다운로드",
                data=docx_bytes or b"",
                file_name="weekly_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        hwpx_bytes = st.session_state.get("report_hwpx")
        with c3:
            st.download_button(
                label="한글(.hwpx) 파일 다운로드",
                data=hwpx_bytes or b"",
                file_name="weekly_report.hwpx",
                mime="application/zip",
            )
        hwpx_err = st.session_state.get("report_hwpx_error", "")
        if hwpx_err:
            st.warning(f"HWPX 생성 경고: {hwpx_err}")


if __name__ == "__main__":
    main()
